"""Генерация HTML/DOCX отчётов: ответы, метрики, графики (HTML)."""
from __future__ import annotations

import io
import json
from html import escape
from typing import Any

from docx import Document
from docxtpl import DocxTemplate

from app.services.test_config import get_test_config, iter_questions


def _metrics_list(metrics: dict | None) -> list[dict[str, Any]]:
    if not metrics:
        return []
    out = []
    for _fid, m in metrics.items():
        if isinstance(m, dict) and "name" in m and "value" in m:
            out.append({"name": m["name"], "value": m["value"]})
    return out


def build_html_report(
    test: Any,
    submission: Any,
    report_kind: str,
) -> str:
    """report_kind: client | specialist"""
    cfg = get_test_config(test)
    label = "Клиент" if report_kind == "client" else "Профориентолог"
    answers = submission.answers or {}
    if isinstance(answers, list):
        answers = {str(i): v for i, v in enumerate(answers)}

    rows = []
    for q in iter_questions(cfg):
        qid = q["id"]
        val = answers.get(qid)
        rows.append(
            f"<tr><td>{escape(q.get('text',''))}</td><td><pre class='mono'>{escape(json.dumps(val, ensure_ascii=False, default=str))}</pre></td></tr>"
        )

    mlist = _metrics_list(submission.metrics)
    metrics_rows = "".join(
        f"<tr><td>{escape(str(m['name']))}</td><td><strong>{escape(str(m['value']))}</strong></td></tr>"
        for m in mlist
    )
    chart_labels = json.dumps([m["name"] for m in mlist], ensure_ascii=False)
    chart_values = json.dumps([float(m["value"]) for m in mlist], ensure_ascii=False)

    return f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8"><title>Отчёт · {escape(test.title)}</title>
<style>
body {{ font-family: system-ui, sans-serif; margin: 2rem; background: #f8fafc; color: #0f172a; }}
h1 {{ color: #1e3a5f; }}
table {{ border-collapse: collapse; width: 100%; background: #fff; margin: 1rem 0; }}
th, td {{ border: 1px solid #cbd5e1; padding: 0.5rem 0.75rem; text-align: left; vertical-align: top; }}
th {{ background: #e2e8f0; }}
pre.mono {{ white-space: pre-wrap; margin: 0; font-size: 0.9rem; }}
.chart-wrap {{ max-width: 720px; margin: 2rem 0; background: #fff; padding: 1rem; border-radius: 8px; border: 1px solid #e2e8f0; }}
</style>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js"></script>
</head><body>
<h1>{escape(test.title)}</h1>
<p><strong>Тип отчёта:</strong> {escape(label)}</p>
<p><strong>Клиент:</strong> {escape(submission.client_name)} ({escape(submission.client_email)})</p>
<p><strong>Дата:</strong> {submission.created_at.isoformat()}</p>

<h2>Ответы</h2>
<table><thead><tr><th>Вопрос</th><th>Ответ</th></tr></thead><tbody>{"".join(rows)}</tbody></table>

<h2>Метрики (формулы)</h2>
<table><thead><tr><th>Показатель</th><th>Значение</th></tr></thead><tbody>{metrics_rows or "<tr><td colspan='2'>Нет расчётных метрик</td></tr>"}</tbody></table>

<div class="chart-wrap"><canvas id="ch"></canvas></div>
<script>
const labels = {chart_labels};
const values = {chart_values};
if (labels.length) {{
  new Chart(document.getElementById('ch'), {{
    type: 'bar',
    data: {{ labels, datasets: [{{ label: 'Метрики', data: values, backgroundColor: '#3b82f6' }}] }},
    options: {{ responsive: true, plugins: {{ legend: {{ display: false }} }} }}
  }});
}}
</script>
</body></html>"""


def build_docx_report(
    test: Any,
    submission: Any,
    report_kind: str,
    template_bytes: bytes | None,
) -> io.BytesIO:
    label = "Клиент" if report_kind == "client" else "Профориентолог"
    cfg = get_test_config(test)
    answers = submission.answers or {}
    if isinstance(answers, list):
        answers = {str(i): v for i, v in enumerate(answers)}

    ctx = {
        "report_kind": label,
        "test_title": test.title,
        "client_name": submission.client_name,
        "client_email": submission.client_email,
        "client_phone": submission.client_phone or "",
        "answers": answers,
        "questions": iter_questions(cfg),
        "metrics": submission.metrics or {},
        "metrics_list": _metrics_list(submission.metrics),
    }

    if template_bytes:
        tpl = DocxTemplate(io.BytesIO(template_bytes))
        tpl.render(ctx)
        out = io.BytesIO()
        tpl.save(out)
        out.seek(0)
        return out

    doc = Document()
    doc.add_heading(f"{test.title} — отчёт ({label})", level=0)
    doc.add_paragraph(f"Клиент: {submission.client_name}, {submission.client_email}")
    doc.add_paragraph("Ответы:")
    for q in iter_questions(cfg):
        qid = q["id"]
        doc.add_paragraph(f"{q.get('text','')}: {answers.get(qid)}")
    doc.add_paragraph("Метрики:")
    for m in _metrics_list(submission.metrics):
        doc.add_paragraph(f"{m['name']}: {m['value']}")
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
